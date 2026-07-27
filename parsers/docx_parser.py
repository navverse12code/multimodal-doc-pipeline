import logging
from docx import Document
from parsers.base_parser import BaseParser

logger = logging.getLogger(__name__)

class DocxParser(BaseParser):
    def parse(self, file_path: str):
        content_sections = []
        try:
            doc = Document(file_path)
            full_text = []

            # Extract Paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    full_text.append(p.text.strip())

            # Extract Tables cleanly
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        table_data.append(" | ".join(row_cells))
                if table_data:
                    full_text.append("\n[Table Data]:\n" + "\n".join(table_data))

            combined_text = "\n\n".join(full_text)
            if combined_text.strip():
                content_sections.append((1, combined_text))

        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")

        return content_sections
